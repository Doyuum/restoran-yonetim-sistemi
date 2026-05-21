from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2.5)
    section.bottom_margin = Cm(2.5)
    section.left_margin   = Cm(3)
    section.right_margin  = Cm(2.5)


def cell_yaz(cell, text, bold=False, size=11, align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    r = p.add_run(text)
    r.bold = bold
    r.font.size = Pt(size)


# ── BAŞLIK TABLOSU ──────────────────────────────────────────────
tbl = doc.add_table(rows=5, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

r0 = tbl.rows[0]
r0.cells[0].merge(r0.cells[1])
cell_yaz(r0.cells[0],
    "MESLEK YÜKSEKOKULU\n2025 – 2026 EĞİTİM – ÖĞRETİM YILI / BAHAR DÖNEMİ",
    bold=True, size=11, align=WD_ALIGN_PARAGRAPH.CENTER)

cell_yaz(tbl.rows[1].cells[0], "Adı / Name",           bold=True, size=10)
cell_yaz(tbl.rows[1].cells[1], "PYTHON PROGRAMLAMA",   size=10)
cell_yaz(tbl.rows[2].cells[0], "DERS Kodu / Code",     bold=True, size=10)
cell_yaz(tbl.rows[2].cells[1], "BLP 276",              size=10)
cell_yaz(tbl.rows[3].cells[0], "Sorumlusu / Lecturer", bold=True, size=10)
cell_yaz(tbl.rows[3].cells[1], "Öğr.Gör. İLKER DURAN",size=10)

r4 = tbl.rows[4]
r4.cells[0].merge(r4.cells[1])
cell_yaz(r4.cells[0], "PROJE HAFTALIK RAPORU",
         bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)

doc.add_paragraph()

hk = doc.add_table(rows=2, cols=2)
hk.style = "Table Grid"
cell_yaz(hk.rows[0].cells[0], "HAFTA", bold=True, size=10)
cell_yaz(hk.rows[0].cells[1], "4. Hafta", size=10)
cell_yaz(hk.rows[1].cells[0], "KONU",  bold=True, size=10)
cell_yaz(hk.rows[1].cells[1], "Test, Senaryolar ve Değerlendirme", size=10)

doc.add_paragraph()

bp = doc.add_paragraph()
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = bp.add_run("HAFTALIK YAPILAN İŞLEMLER")
r.bold = True
r.font.size = Pt(11)
r.font.underline = True

doc.add_paragraph()


def yaz(metin, bold=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.space_before = Pt(6) if bold else Pt(0)
    if not bold:
        p.paragraph_format.first_line_indent = Cm(1)
    run = p.add_run(metin)
    run.font.size = Pt(11)
    run.bold = bold


yaz("4.1. Senaryo Tabanlı Testler", bold=True)

yaz(
    "bu hafta sistemi test etmek için birkaç farklı senaryo denedik. "
    "ilk olarak müşteri gibi giriş yapıp qr kod ile masaya sipariş verdik "
    "menüden ürün seçip onaylayınca siparişin mutfak ekranına geçtiğini gördük "
    "sonra mutfaktan hazırlandı yapınca kasa ekranına geçti, bu kısım çalıştı."
)

yaz(
    "ikinci testte malzeme stoğunu sıfırladık ve o malzemeyi kullanan yemeğin "
    "menüde tükendi yazıp yazamadığına baktık. sistem doğru şekilde o yemeği "
    "kapattı. malzeme tekrar eklenince yemek tekrar açıldı."
)

yaz(
    "üçüncü testte dolu masaya tekrar sipariş açmaya çalıştık sistem izin vermedi "
    "ayrıca rezerve masayı seçmeye çalıştık onu da engelledi. bu senaryolar "
    "beklediğimiz gibi sonuçlandı."
)

yaz("4.2. Hataların Giderilmesi", bold=True)

yaz(
    "testler sırasında bazı hatalar fark ettik. en çok dikkatimizi çeken şey "
    "müşteri masa seçip sipariş vermeden çıkınca masanın dolu kalmasıydı "
    "bunu düzeltmek için program kapanırken masaların otomatik boşalması sağlandı."
)

yaz(
    "bir de sipariş numaraları her gün 1 den başlamıyordu bunu fark edince "
    "gün değişimini kontrol eden bir yapı ekledik artık her sabah 1 den başlıyor. "
    "türkçe karakter sorunu da vardı terminalde harfler bozuk çıkıyordu "
    "encoding ayarı yaparak düzelttik."
)

yaz("4.3. Projenin Amaca Uygunluğu", bold=True)

yaz(
    "projeye başlarken müşterilerin sipariş verebilmesi mutfağın bunu görmesi "
    "ve kasanın ödeme alması hedefleniyordu. bu üç şey çalışıyor. "
    "bunlara ek olarak malzeme takibi kampanya sistemi ve günlük rapor da eklendi "
    "yani hedeflediğimizden daha fazlası yapılmış oldu."
)

yaz(
    "menü düzenleme stok güncelleme masa rezervasyonu gibi yönetim işlemleri de "
    "var. veriler dosyaya kaydedildiği için program kapansa bile bilgiler silinmiyor."
)

yaz("4.4. Gerçek Hayat Uyumu", bold=True)

yaz(
    "sistemin gerçek bir restoranda kullanılabileceğini düşünüyoruz "
    "müşteriler masadan qr okutup telefon üzerinden sipariş verebildiği için "
    "garson gereksinimi azalıyor. mutfak sırayla siparişleri görüyor "
    "kasa ödeme ve para üstü hesabını kolayca yapabiliyor."
)

yaz(
    "tek şart restoranın wifi ağı olması, internet gerekmyor sadece yerel ağ yeterli "
    "bu da sistemi daha kullanışlı yapıyor. "
    "dört hafta içinde sıfırdan başlayıp kullanılabilir bir sistem çıkarttık "
    "test sürecinde de büyük sorun yaşamadık."
)

# ── İMZA TABLOSU ────────────────────────────────────────────────
doc.add_page_break()

it = doc.add_table(rows=4, cols=3)
it.style = "Table Grid"
it.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, b in enumerate(["NUMARA", "GRUP ÜYELERİNİN İSİMLERİ", "İMZA"]):
    cell_yaz(it.rows[0].cells[i], b, bold=True, size=10,
             align=WD_ALIGN_PARAGRAPH.CENTER)

uyeler = [
    ("202407124023", "Ege Recep Alembeyli"),
    ("202407124030", "Taha Saranbeyli"),
]
for idx, (numara, isim) in enumerate(uyeler, 1):
    cell_yaz(it.rows[idx].cells[0], numara, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_yaz(it.rows[idx].cells[1], isim, size=10)
    it.rows[idx].cells[2].text = ""

cell_yaz(it.rows[3].cells[0], "", size=10)
cell_yaz(it.rows[3].cells[1], "Öğr.Gör. İlker Duran", bold=True, size=10)
cell_yaz(it.rows[3].cells[2], "İmza", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

kayit = r"C:\Users\Doyuum\Desktop\BLP_Grup11_4Hafta_Rapor.docx"
doc.save(kayit)
print("Kaydedildi:", kayit)
