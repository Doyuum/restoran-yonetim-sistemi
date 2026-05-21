"""
Restoran Yönetim Sistemi — Word Raporu Oluşturucu
Çalıştır: python rapor_olustur.py
"""

from docx import Document
from docx.shared import Pt, RGBColor, Cm, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import datetime

# ── Renkler ─────────────────────────────────────────────────────
KOYU_LACIVERT   = "0D1B2A"
ORTA_LACIVERT   = "1B2A4A"
MAVI            = "1F4E79"
ACIK_MAVI       = "2E74B5"
TEAL            = "0F6674"
ACIK_TEAL       = "D6F0F5"
GECE_MAVISI     = "162D40"
BEYAZ           = "FFFFFF"
ACIK_GRI        = "F2F2F2"
KOYU_METIN      = "1A1A2E"
ALTIN           = "C8A951"

# ── Yardımcı: hücre arka planı ──────────────────────────────────
def hucre_renk(hucre, hex_renk):
    tc   = hucre._tc
    tcPr = tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:fill"),  hex_renk)
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:val"),   "clear")
    tcPr.append(shd)

# ── Yardımcı: satır yüksekliği ──────────────────────────────────
def satir_yukseklik(satir, yukseklik_cm):
    tr   = satir._tr
    trPr = tr.get_or_add_trPr()
    trH  = OxmlElement("w:trHeight")
    trH.set(qn("w:val"),   str(int(yukseklik_cm * 567)))
    trH.set(qn("w:hRule"), "exact")
    trPr.append(trH)

# ── Yardımcı: tablo kenarlık kaldır ─────────────────────────────
def kenarlık_kaldir(tablo):
    tbl  = tablo._tbl
    tblPr = tbl.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr")
        tbl.insert(0, tblPr)
    borders = OxmlElement("w:tblBorders")
    for kenar in ("top","left","bottom","right","insideH","insideV"):
        el = OxmlElement(f"w:{kenar}")
        el.set(qn("w:val"),   "none")
        el.set(qn("w:sz"),    "0")
        el.set(qn("w:space"), "0")
        el.set(qn("w:color"), "auto")
        borders.append(el)
    tblPr.append(borders)

# ── Yardımcı: çalışma sayfası arka plan rengi ───────────────────
def sayfa_arka_plan(bolum, hex_renk):
    sectPr = bolum._sectPr
    bg = OxmlElement("w:background")
    bg.set(qn("w:color"), hex_renk)
    doc.element.insert(0, bg)

# ── Yardımcı: metin ekle (hücre içine) ─────────────────────────
def hucre_yaz(hucre, metin, boyut=11, kalin=False, italik=False,
              renk=KOYU_METIN, hizalama=WD_ALIGN_PARAGRAPH.LEFT,
              aralik_once=0, aralik_sonra=0):
    p = hucre.paragraphs[0] if hucre.paragraphs else hucre.add_paragraph()
    p.alignment = hizalama
    p.paragraph_format.space_before = Pt(aralik_once)
    p.paragraph_format.space_after  = Pt(aralik_sonra)
    run = p.add_run(metin)
    run.bold   = kalin
    run.italic = italik
    run.font.size  = Pt(boyut)
    run.font.color.rgb = RGBColor.from_string(renk)
    return p

def hucre_ekle_paragraf(hucre, metin, boyut=11, kalin=False, italik=False,
                         renk=KOYU_METIN, hizalama=WD_ALIGN_PARAGRAPH.LEFT,
                         aralik_once=4, aralik_sonra=0):
    p = hucre.add_paragraph()
    p.alignment = hizalama
    p.paragraph_format.space_before = Pt(aralik_once)
    p.paragraph_format.space_after  = Pt(aralik_sonra)
    run = p.add_run(metin)
    run.bold   = kalin
    run.italic = italik
    run.font.size  = Pt(boyut)
    run.font.color.rgb = RGBColor.from_string(renk)
    return p

# ════════════════════════════════════════════════════════════════
doc = Document()

# Sayfa ayarları (normal kenar boşlukları, kapak hariç)
for bolum in doc.sections:
    bolum.page_width    = Cm(21)
    bolum.page_height   = Cm(29.7)
    bolum.top_margin    = Cm(2.0)
    bolum.bottom_margin = Cm(2.0)
    bolum.left_margin   = Cm(2.5)
    bolum.right_margin  = Cm(2.5)

# ════════════════════════════════════════════════════════════════
#  KAPAK SAYFASI
# ════════════════════════════════════════════════════════════════
# Tam sayfa lacivert tablo
kapak = doc.add_table(rows=4, cols=1)
kenarlık_kaldir(kapak)
kapak.alignment = WD_TABLE_ALIGNMENT.CENTER
kapak.columns[0].width = Cm(21)

# Satır yükseklikleri
satir_yukseklik(kapak.rows[0], 5.0)   # üst boşluk
satir_yukseklik(kapak.rows[1], 6.0)   # logo & başlık
satir_yukseklik(kapak.rows[2], 10.0)  # içerik
satir_yukseklik(kapak.rows[3], 5.0)   # alt

for satir in kapak.rows:
    hucre_renk(satir.cells[0], KOYU_LACIVERT)

# Üst boşluk
kapak.rows[0].cells[0].paragraphs[0].clear()

# Logo alanı — simgesel metin
logo_h = kapak.rows[1].cells[0]
hucre_renk(logo_h, KOYU_LACIVERT)
logo_h.paragraphs[0].clear()
logo_h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
p_logo = logo_h.add_paragraph()
p_logo.alignment = WD_ALIGN_PARAGRAPH.CENTER
r = p_logo.add_run("⬡ R · S")
r.bold = True
r.font.size  = Pt(36)
r.font.color.rgb = RGBColor.from_string(TEAL)

p_bas = logo_h.add_paragraph()
p_bas.alignment = WD_ALIGN_PARAGRAPH.CENTER
r2 = p_bas.add_run("RestoranSys")
r2.bold = True
r2.font.size  = Pt(40)
r2.font.color.rgb = RGBColor.from_string(BEYAZ)

# Orta içerik
icerik_h = kapak.rows[2].cells[0]
hucre_renk(icerik_h, KOYU_LACIVERT)
icerik_h.paragraphs[0].clear()
icerik_h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

p_alt = icerik_h.add_paragraph()
p_alt.alignment = WD_ALIGN_PARAGRAPH.LEFT
p_alt.paragraph_format.left_indent = Cm(2)
r3 = p_alt.add_run(f"2026 RestoranSys Raporu\n")
r3.font.size  = Pt(14)
r3.font.color.rgb = RGBColor.from_string(BEYAZ)
r4 = p_alt.add_run("Python Programlama")
r4.bold = True
r4.font.size  = Pt(14)
r4.font.color.rgb = RGBColor.from_string(BEYAZ)

p_slogan = icerik_h.add_paragraph()
p_slogan.alignment = WD_ALIGN_PARAGRAPH.CENTER
p_slogan.paragraph_format.space_before = Pt(30)
r5 = p_slogan.add_run("Restoranınızın verimliliği için RestoranSys yanınızda")
r5.italic = True
r5.font.size  = Pt(18)
r5.font.color.rgb = RGBColor.from_string(BEYAZ)

# Alt — isimler
alt_h = kapak.rows[3].cells[0]
hucre_renk(alt_h, GECE_MAVISI)
alt_h.paragraphs[0].clear()
alt_h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

p_isim = alt_h.add_paragraph()
p_isim.alignment = WD_ALIGN_PARAGRAPH.CENTER
r6 = p_isim.add_run("Tarafından Hazırlanmıştır.")
r6.font.size  = Pt(10)
r6.font.color.rgb = RGBColor.from_string(BEYAZ)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  İÇİNDEKİLER SAYFASI
# ════════════════════════════════════════════════════════════════

# Üst dekoratif bant
ust_bant = doc.add_table(rows=1, cols=1)
kenarlık_kaldir(ust_bant)
satir_yukseklik(ust_bant.rows[0], 1.8)
hucre_renk(ust_bant.rows[0].cells[0], TEAL)
h = ust_bant.rows[0].cells[0]
h.paragraphs[0].clear()
h.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
pb = h.add_paragraph()
pb.alignment = WD_ALIGN_PARAGRAPH.LEFT
rb = pb.add_run("   📋  Rapor İçerikleri")
rb.bold = True
rb.font.size  = Pt(16)
rb.font.color.rgb = RGBColor.from_string(BEYAZ)

doc.add_paragraph()

# İçindekiler başlık
p_ic = doc.add_paragraph()
p_ic.alignment = WD_ALIGN_PARAGRAPH.LEFT
ric = p_ic.add_run("Rapor İçerikleri")
ric.bold = True
ric.font.size  = Pt(28)
ric.font.color.rgb = RGBColor(0x0D, 0x1B, 0x2A)

doc.add_paragraph()

# 2 sütunlu içindekiler tablosu
ic_tablo = doc.add_table(rows=2, cols=2)
kenarlık_kaldir(ic_tablo)

ic_satirlar = [
    ("03", "Proje Amacının Tanımlanması"),
    ("04", "Proje Türü ve Kapsamı"),
    ("05", "Sistem Bileşenleri ve Akış"),
    ("06", "Temel Özelliklerin Belirlenmesi"),
]

for i, (no, baslik_ic) in enumerate(ic_satirlar):
    satir_idx = i // 2
    sutun_idx = i %  2
    h2 = ic_tablo.rows[satir_idx].cells[sutun_idx]
    hucre_renk(h2, ACIK_TEAL if i % 2 == 0 else ACIK_GRI)
    h2.paragraphs[0].clear()
    h2.vertical_alignment = WD_ALIGN_VERTICAL.CENTER

    p_no = h2.add_paragraph()
    p_no.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_no.paragraph_format.space_before = Pt(10)
    r_no = p_no.add_run(no)
    r_no.bold = True
    r_no.font.size  = Pt(40)
    r_no.font.color.rgb = RGBColor.from_string(TEAL)

    p_bc = h2.add_paragraph()
    p_bc.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_bc.paragraph_format.space_after = Pt(10)
    r_bc = p_bc.add_run(baslik_ic)
    r_bc.bold = True
    r_bc.font.size  = Pt(12)
    r_bc.font.color.rgb = RGBColor.from_string(MAVI)

# Alt bant (sayfa no)
doc.add_paragraph()
alt_bant = doc.add_table(rows=1, cols=3)
kenarlık_kaldir(alt_bant)
satir_yukseklik(alt_bant.rows[0], 0.8)
for c in alt_bant.rows[0].cells:
    hucre_renk(c, ORTA_LACIVERT)
alt_bant.rows[0].cells[0].paragraphs[0].clear()
alt_bant.rows[0].cells[1].paragraphs[0].clear()
alt_bant.rows[0].cells[2].paragraphs[0].clear()

p_sol = alt_bant.rows[0].cells[0].add_paragraph()
p_sol.alignment = WD_ALIGN_PARAGRAPH.LEFT
r_sol = p_sol.add_run("  1. HAFTA RAPORU")
r_sol.font.size = Pt(9)
r_sol.font.color.rgb = RGBColor.from_string(BEYAZ)

p_orta = alt_bant.rows[0].cells[1].add_paragraph()
p_orta.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_orta = p_orta.add_run("02")
r_orta.bold = True
r_orta.font.size = Pt(9)
r_orta.font.color.rgb = RGBColor.from_string(BEYAZ)

p_sag = alt_bant.rows[0].cells[2].add_paragraph()
p_sag.alignment = WD_ALIGN_PARAGRAPH.RIGHT
r_sag = p_sag.add_run(f"NİSAN 2026  ")
r_sag.font.size = Pt(9)
r_sag.font.color.rgb = RGBColor.from_string(BEYAZ)

doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SAYFA ŞABLONU (footer bant + içerik)
# ════════════════════════════════════════════════════════════════

def alt_bant_ekle(sayfa_no):
    doc.add_paragraph()
    bant = doc.add_table(rows=1, cols=3)
    kenarlık_kaldir(bant)
    satir_yukseklik(bant.rows[0], 0.8)
    for c in bant.rows[0].cells:
        hucre_renk(c, ORTA_LACIVERT)
    for p in [c.paragraphs[0] for c in bant.rows[0].cells]:
        p.clear()
    s = bant.rows[0].cells[0].add_paragraph()
    s.alignment = WD_ALIGN_PARAGRAPH.LEFT
    rs = s.add_run("  1. HAFTA RAPORU")
    rs.font.size = Pt(9)
    rs.font.color.rgb = RGBColor.from_string(BEYAZ)
    m = bant.rows[0].cells[1].add_paragraph()
    m.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rm = m.add_run(str(sayfa_no).zfill(2))
    rm.bold = True
    rm.font.size = Pt(9)
    rm.font.color.rgb = RGBColor.from_string(BEYAZ)
    sa = bant.rows[0].cells[2].add_paragraph()
    sa.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rsa = sa.add_run("NİSAN 2026  ")
    rsa.font.size = Pt(9)
    rsa.font.color.rgb = RGBColor.from_string(BEYAZ)

def bolum_baslik(no, metin):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after  = Pt(8)
    r = p.add_run(f"{no}  {metin}")
    r.bold = True
    r.font.size  = Pt(22)
    r.font.color.rgb = RGBColor.from_string(MAVI)
    return p

def alt_baslik(metin):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(metin)
    r.bold = True
    r.font.size  = Pt(13)
    r.font.color.rgb = RGBColor.from_string(TEAL)
    return p

def metin_ekle(icerik, boyut=10.5, italik=False):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    r = p.add_run(icerik)
    r.italic = italik
    r.font.size  = Pt(boyut)
    r.font.color.rgb = RGBColor.from_string(KOYU_METIN)
    return p

def ozellik_ekle(numara, baslik_oz, aciklama):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after  = Pt(2)
    p.paragraph_format.left_indent  = Cm(0.3)
    r1 = p.add_run(f"{numara}. {baslik_oz}\n")
    r1.bold = True
    r1.font.size  = Pt(11)
    r1.font.color.rgb = RGBColor.from_string(ACIK_MAVI)
    r2 = p.add_run(aciklama)
    r2.font.size  = Pt(10)
    r2.font.color.rgb = RGBColor.from_string(KOYU_METIN)

# ════════════════════════════════════════════════════════════════
#  SAYFA 3 — 1.1 Proje Amacının Tanımlanması
# ════════════════════════════════════════════════════════════════
bolum_baslik("1.1", "Proje Amacının Tanımlanması")

# İki sütunlu tablo
iki_sutun = doc.add_table(rows=1, cols=2)
kenarlık_kaldir(iki_sutun)
iki_sutun.columns[0].width = Cm(8.5)
iki_sutun.columns[1].width = Cm(8.5)

sol = iki_sutun.rows[0].cells[0]
sag = iki_sutun.rows[0].cells[1]
sol.paragraphs[0].clear()
sag.paragraphs[0].clear()

hucre_yaz(sol, "Projenin İçeriği", boyut=12, kalin=True, renk=TEAL,
          aralik_sonra=6)
hucre_ekle_paragraf(sol,
    "Bu proje, restoranların günlük operasyonlarını tek bir platform "
    "üzerinden dijital olarak yönetmesine olanak tanıyan çok terminalli "
    "bir Python konsol uygulamasıdır. Trello ve benzeri yönetim araçlarından "
    "ilham alınarak geliştirilmiş bu sistem; sipariş takibi, masa yönetimi, "
    "mutfak koordinasyonu ve kasa işlemlerini birbirinden bağımsız terminaller "
    "aracılığıyla sunar.",
    boyut=10, renk=KOYU_METIN)
hucre_ekle_paragraf(sol,
    "Sistem temel olarak şu işlemleri gerçekleştirmektedir: Müşteriler "
    "sipariş ekranından masa seçerek ürün siparişi verebilmektedir. "
    "Mutfak terminali gelen siparişleri anlık olarak görmekte ve durum "
    "güncellemesi yapabilmektedir. Kasa terminali ödemeleri almakta ve "
    "para üstü hesaplamaktadır. Yönetim terminali ise menü, stok, kampanya "
    "ve malzeme yönetimini kapsamaktadır.",
    boyut=10, renk=KOYU_METIN)

hucre_yaz(sag, "Hangi Kullanıcı İhtiyacını Karşılamaktadır",
          boyut=12, kalin=True, renk=TEAL, aralik_sonra=6)
hucre_ekle_paragraf(sag,
    "Günümüzde küçük ve orta ölçekli restoranlar sipariş yönetimini "
    "çoğunlukla kağıt kalem ile ya da mesajlaşma uygulamaları üzerinden "
    "yürütmektedir. Bu durum sipariş karışıklığına, mutfak ile kasa arasındaki "
    "koordinasyon eksikliğine ve stok takibinin güçleşmesine yol açmaktadır.",
    boyut=10, renk=KOYU_METIN)
hucre_ekle_paragraf(sag,
    "RestoranSys; müşteriden mutfağa, garsondan kasaya uzanan tüm süreci "
    "tek bir çatı altında dijitalleştirerek bu sorunları çözmektedir. "
    "Anlık stok takibi, otomatik malzeme düşme sistemi ve kampanya/indirim "
    "modülü ile restoranın hem operasyonel hem de ticari verimliliğini "
    "artırmaktadır.",
    boyut=10, renk=KOYU_METIN)

alt_bant_ekle(3)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SAYFA 4 — 1.2 Proje Türü ve Kapsamı
# ════════════════════════════════════════════════════════════════
bolum_baslik("1.2", "Proje Türü ve Kapsamı")

metin_ekle(
    "Geliştirilen uygulama, tamamen Python 3 ile yazılmış terminal tabanlı "
    "çok katmanlı bir yönetim sistemidir. Herhangi bir ek web sunucusuna veya "
    "tarayıcıya ihtiyaç duymaksızın komut satırı üzerinden doğrudan çalıştırılabilmektedir. "
    "Veriler JSON dosyalarında saklanmakta; böylece hafif, hızlı ve kurulum gerektirmeyen "
    "bir yapı elde edilmektedir."
)
metin_ekle(
    "Uygulama; restoran sahipleri, mutfak personeli, garsonlar ve kasiyer olmak üzere "
    "farklı rol gruplarını hedef kitle olarak benimsemektedir. Proje kapsamında dört temel "
    "terminal modülü ele alınmaktadır: Müşteri Terminali, Mutfak Terminali, Kasa Terminali "
    "ve Yönetim Terminali. Her terminal bağımsız çalışmakla birlikte aynı veri tabanını "
    "paylaşmaktadır."
)

# Renkli bilgi kutusu
bilgi_t = doc.add_table(rows=1, cols=1)
kenarlık_kaldir(bilgi_t)
hucre_renk(bilgi_t.rows[0].cells[0], ACIK_TEAL)
bilgi_h = bilgi_t.rows[0].cells[0]
bilgi_h.paragraphs[0].clear()
for satir_m in [
    "  Kullanılan Teknolojiler:",
    "  • Python 3.10+   • msvcrt (gerçek zamanlı klavye girişi)",
    "  • json (veri kalıcılığı)   • dataclasses & enum (veri modelleri)",
    "  • python-docx (rapor oluşturma)   • pathlib (dosya yönetimi)",
]:
    p_b = bilgi_h.add_paragraph()
    r_b = p_b.add_run(satir_m)
    r_b.font.size  = Pt(10)
    r_b.font.color.rgb = RGBColor.from_string(MAVI)
    r_b.bold = "Teknolojiler" in satir_m

alt_bant_ekle(4)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SAYFA 5 — 1.3 Sistem Bileşenleri ve Akış
# ════════════════════════════════════════════════════════════════
bolum_baslik("1.3", "Sistem Bileşenleri ve Akış")

metin_ekle("Sistem dört ana bileşenden oluşmaktadır", italik=True)

# 2×2 bileşen tablosu
bilesen_t = doc.add_table(rows=2, cols=2)
kenarlık_kaldir(bilesen_t)

bilesenler = [
    ("TERMINAL BİLEŞENİ",
     "Sisteme erişim, terminal seçim ekranıyla başlamaktadır. "
     "Müşteri terminali şifresiz açılırken Mutfak, Kasa ve Yönetim "
     "terminalleri şifre korumalıdır. Her terminal kendi yetki sınırları "
     "içinde işlem yapabilmekte; aynı JSON veri tabanı üzerinden eş zamanlı "
     "çalışabilmektedir. 20 saniyelik hareketsizlik zaman aşımı müşteri "
     "ekranlarını otomatik olarak giriş sayfasına döndürür."),
    ("MENÜ BİLEŞENİ",
     "Oturum açan yönetici; menüye ürün ekleyebilmekte, fiyat ve stok "
     "güncelleyebilmekte, kampanya/indirim uygulayabilmektedir. Her ürün "
     "için stok modu tanımlanabilir: sınırsız (-1), sayılı (>0) veya "
     "tükenmiş (0). QR komutu ile herhangi bir ekranda anlık menü raporu "
     "görüntülenebilmektedir."),
    ("SİPARİŞ BİLEŞENİ",
     "Proje içerisinde siparişler oluşturulabilmekte, düzenlenebilmekte "
     "ve durumları güncellenebilmektedir. Her sipariş; masa numarası, "
     "kalemler, KDV dahil toplam tutar ve sipariş notu bilgilerini "
     "barındırmaktadır. Siparişin durumu altı aşamalı bir durum makinesiyle "
     "takip edilmektedir."),
    ("VERİ BİLEŞENİ",
     "Tüm menü, sipariş, masa ve malzeme verileri JSON dosyalarında kalıcı "
     "olarak saklanmaktadır. Sistem her yeniden başlatıldığında mevcut veriler "
     "yüklenerek kullanıcıya sunulmaktadır. Mutfak slibi her sipariş için "
     "veri/slipler/ klasörüne .txt dosyası olarak otomatik kaydedilmektedir."),
]

for i, (bas, acik) in enumerate(bilesenler):
    satir_i = i // 2
    sutun_i = i %  2
    h3 = bilesen_t.rows[satir_i].cells[sutun_i]
    hucre_renk(h3, ACIK_TEAL if i % 2 == 0 else ACIK_GRI)
    h3.paragraphs[0].clear()
    h3.vertical_alignment = WD_ALIGN_VERTICAL.TOP
    hucre_yaz(h3, bas, boyut=11, kalin=True, renk=TEAL,
              aralik_once=6, aralik_sonra=4)
    hucre_ekle_paragraf(h3, acik, boyut=9.5, renk=KOYU_METIN, aralik_once=2)

# Sistem akış kutusu
doc.add_paragraph()
akis_t = doc.add_table(rows=1, cols=1)
kenarlık_kaldir(akis_t)
hucre_renk(akis_t.rows[0].cells[0], ORTA_LACIVERT)
akis_h = akis_t.rows[0].cells[0]
akis_h.paragraphs[0].clear()
p_akis = akis_h.add_paragraph()
p_akis.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_akis = p_akis.add_run(
    "Sistem akışı:  "
    "Müşteri girişi  →  Masa seç  →  Sipariş ver  →  "
    "Mutfak onayı  →  Garson servisi  →  Kasa ödemesi"
)
r_akis.italic = True
r_akis.font.size  = Pt(10)
r_akis.font.color.rgb = RGBColor.from_string(BEYAZ)

alt_bant_ekle(5)
doc.add_page_break()

# ════════════════════════════════════════════════════════════════
#  SAYFA 6 — 1.4 Temel Özelliklerin Belirlenmesi
# ════════════════════════════════════════════════════════════════
bolum_baslik("1.4", "Temel Özelliklerin Belirlenmesi")

ozellikler = [
    ("Çok Terminalli Mimari",
     "Uygulama Müşteri, Mutfak, Kasa ve Yönetim olmak üzere dört bağımsız "
     "terminal moduna sahiptir. Her terminal farklı yetki seviyesinde çalışır; "
     "Müşteri terminali şifresizken diğerleri şifre korumalıdır."),
    ("Stok ve Malzeme Yönetimi",
     "Menü ürünlerinin porsiyon stoğu anlık takip edilmektedir. Çorba "
     "ürünleri için ayrı bir malzeme stok sistemi bulunmakta; sipariş "
     "verildiğinde tarife göre malzemeler otomatik düşülmektedir."),
    ("Kampanya ve İndirim Sistemi",
     "Yönetim terminali üzerinden herhangi bir ürüne yüzde tabanlı indirim "
     "uygulanabilmektedir. İndirimli ürünler müşteri ekranında orijinal fiyat "
     "ve kampanya yüzdesiyle birlikte gösterilmektedir."),
    ("KDV Dahil Fiyatlandırma",
     "Müşteri ekranında tüm fiyatlar KDV dahil (%10 KDV) olarak "
     "gösterilmektedir. Personel ekranlarında ise KDV hariç taban fiyat "
     "görüntülenir; KDV ayrıca hesaplanarak faturada belirtilmektedir."),
    ("QR Anlık Menü Raporu",
     "Herhangi bir ekranda 'QR' yazıldığında ürün listesi, stok durumu ve "
     "kampanya bilgilerini içeren tam menü raporu anında görüntülenmektedir. "
     "Bu özellik hem müşteri hem de personel için kullanılabilmektedir."),
    ("Masa Rezervasyon Sistemi",
     "Personel, Yönetim Terminali üzerinden masaları Boş, Dolu veya Rezerve "
     "olarak işaretleyebilmektedir. Rezerve masalar müşteri tarafından "
     "seçilememektedir."),
    ("Otomatik Zaman Aşımı ve Güvenlik",
     "Müşteri ekranında 20 saniye hareketsizlik durumunda sipariş otomatik "
     "iptal edilir ve masa boşaltılır. Şifre girerken karakterler '*' olarak "
     "maskelenmekte; 4 hatalı denemeden sonra giriş engellenmektedir."),
]

for i, (bas, acik) in enumerate(ozellikler, 1):
    ozellik_ekle(i, bas, acik)

alt_bant_ekle(6)

# ── Kaydet ──────────────────────────────────────────────────────
cikti = r"C:\Users\Doyuum\Desktop\Restoran_Yonetim_Sistemi_Rapor.docx"
doc.save(cikti)
print(f"Rapor olusturuldu: {cikti}")
