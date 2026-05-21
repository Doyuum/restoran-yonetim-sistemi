from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

doc = Document()

for section in doc.sections:
    section.top_margin    = Cm(2)
    section.bottom_margin = Cm(2)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def cell_font(cell, text, bold=False, size=11, color=None,
              align=WD_ALIGN_PARAGRAPH.LEFT):
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = RGBColor.from_string(color)


# ── BAŞLIK TABLOSU ──────────────────────────────────────────────
tbl = doc.add_table(rows=5, cols=2)
tbl.style = "Table Grid"
tbl.alignment = WD_TABLE_ALIGNMENT.CENTER

for row in tbl.rows:
    row.cells[0].width = Cm(5)
    row.cells[1].width = Cm(12.5)

# Satır 0 – okul adı (birleşik)
r0 = tbl.rows[0]
r0.cells[0].merge(r0.cells[1])
cell_font(r0.cells[0],
          "MESLEK YÜKSEKOKULU\n2025 – 2026 EĞİTİM – ÖĞRETİM YILI / BAHAR DÖNEMİ",
          bold=True, size=12, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r0.cells[0], "1F3864")
r0.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

cell_font(tbl.rows[1].cells[0], "Adı / Name",          bold=True, size=10)
cell_font(tbl.rows[1].cells[1], "PYTHON PROGRAMLAMA",  size=10)
cell_font(tbl.rows[2].cells[0], "DERS Kodu / Code",    bold=True, size=10)
cell_font(tbl.rows[2].cells[1], "BLP 276",             size=10)
cell_font(tbl.rows[3].cells[0], "Sorumlusu / Lecturer",bold=True, size=10)
cell_font(tbl.rows[3].cells[1], "Öğr.Gör. İLKER DURAN", size=10)

r4 = tbl.rows[4]
r4.cells[0].merge(r4.cells[1])
cell_font(r4.cells[0], "PROJE HAFTALIK RAPORU",
          bold=True, size=13, align=WD_ALIGN_PARAGRAPH.CENTER)
set_cell_bg(r4.cells[0], "2E75B6")
r4.cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

doc.add_paragraph()

# ── HAFTA / KONU ────────────────────────────────────────────────
hk = doc.add_table(rows=2, cols=2)
hk.style = "Table Grid"
cell_font(hk.rows[0].cells[0], "HAFTA", bold=True, size=10)
cell_font(hk.rows[0].cells[1], "3. Hafta", size=10)
cell_font(hk.rows[1].cells[0], "KONU",  bold=True, size=10)
cell_font(hk.rows[1].cells[1], "Gelişmiş Özellikler ve Proje Zenginleştirme", size=10)

doc.add_paragraph()

# ── HAFTALIK YAPILAN İŞLEMLER ───────────────────────────────────
bp = doc.add_paragraph()
bp.alignment = WD_ALIGN_PARAGRAPH.CENTER
br = bp.add_run("HAFTALIK YAPILAN İŞLEMLER")
br.bold = True
br.font.size = Pt(12)
br.font.underline = True

doc.add_paragraph()

PARAGRAFLAR = [
    ("3.1. Projeye Özgü Gelişmiş Özellikler", True),
    (
        "Bu hafta Restoran Yönetim Sistemi projesine kapsamlı gelişmiş özellikler eklenmiştir. "
        "Projeye özgü en önemli yenilik, QR kod tabanlı gerçek zamanlı sipariş sistemidir. "
        "Müşteriler masalarındaki QR kodu telefon kameralarıyla okutarak kendi cihazlarından "
        "sipariş verebilmekte; bu siparişler anlık olarak mutfak ve kasa terminallerine iletilmektedir. "
        "Flask web çerçevesi kullanılarak geliştirilen bu sistem, aynı WiFi ağındaki tüm cihazlardan "
        "erişilebilir hale getirilmiştir.",
        False,
    ),
    (
        "Sisteme filtreleme ve öneri özellikleri de eklenmiştir. Web menü arayüzünde Başlangıç, "
        "Ana Yemek, Pizza, Salata, Tatlı ve İçecek kategorilerine ait filtre butonları oluşturulmuştur. "
        "Müşteriler bu butonlara tıkladığında yalnızca ilgili kategorinin ürünleri görüntülenmekte, "
        "böylece menüde gezinme deneyimi önemli ölçüde kolaylaşmaktadır.",
        False,
    ),
    (
        "İstatistik ve analiz kapsamında günlük rapor sistemi geliştirilmiştir. Yönetim terminalinden "
        "erişilebilen bu rapor; günlük sipariş sayısını, toplam cirosunu ve en çok sipariş edilen "
        "ürünlerin listesini sunmaktadır. Kampanya ve indirim yönetimi modülü sayesinde ürünlere "
        "yüzdelik indirim uygulanabilmekte, müşteriye orijinal fiyat ile indirimli fiyat birlikte "
        "gösterilmektedir.",
        False,
    ),
    ("3.2. Kullanıcı Deneyimi (UI/UX)", True),
    (
        "Web arayüzü tamamen mobil uyumlu olarak tasarlanmıştır. Her menü ürünü kendi kartında "
        "gösterilmekte; ürün adı, açıklaması, fiyatı (KDV dahil) ve sepete ekleme butonu tek bir "
        "bakışta algılanabilecek şekilde düzenlenmiştir. Kampanyalı ürünlerde eski fiyat üzeri çizili "
        "olarak gösterilmekte, indirim oranı renkli rozet ile vurgulanmaktadır.",
        False,
    ),
    (
        "Menü yapısı düzenli ve sezgisel tutulmuştur. Terminal ekranlarında her işlem için klavye "
        "kısayolları açıkça belirtilmekte, hata durumlarında kullanıcıya anlaşılır geri bildirimler "
        "verilmektedir. Müşteri siparişini onayladıktan sonra sipariş durumunu takip edebildiği "
        "bir durum sayfasına yönlendirilmektedir.",
        False,
    ),
    ("3.3. Modüler Kod Yapısı", True),
    (
        "Proje başından itibaren modüler bir mimari ile geliştirilmiştir. Her işlev kendi dosyasında "
        "ayrı bir modül olarak tanımlanmıştır: models.py veri modellerini, storage.py JSON okuma/yazma "
        "işlemlerini, menu_manager.py menü yönetimini, order_manager.py sipariş yönetimini, "
        "malzeme_manager.py malzeme stok takibini ve web_sunucu.py QR sipariş web sunucusunu "
        "kapsamaktadır.",
        False,
    ),
    (
        "Nesne yönelimli programlama (OOP) ilkeleri tercih edilmiştir. MenuYoneticisi, "
        "SiparisYoneticisi ve MalzemeYoneticisi sınıfları, kendi sorumluluk alanlarını "
        "kapsülleyerek bağımsız birimler halinde çalışmaktadır. Bu yapı sayesinde yeni özellik "
        "eklemek ve mevcut kodu test etmek kolaylaşmaktadır.",
        False,
    ),
    ("3.4. Performans ve Temizlik", True),
    (
        "Kodun okunabilirliği ve sürdürülebilirliği ön planda tutulmuştur. Tekrar eden kod blokları "
        "fonksiyonlara taşınmış, her fonksiyon tek bir görevi yerine getirmektedir. JSON tabanlı "
        "kalıcı depolama sistemi sayesinde veriler hızlı okunup yazılmakta, gereksiz dosya işlemleri "
        "önlenmektedir.",
        False,
    ),
    (
        "Performans açısından web sunucusu arka planda ayrı bir thread üzerinde çalıştırılmakta, "
        "terminal arayüzü ile web arayüzü birbirini bloke etmemektedir. Mutfak ve kasa terminalleri "
        "QR ile gelen yeni siparişleri görebilmek için her döngüde JSON verilerini yenilemektedir. "
        "Bu yaklaşım gereksiz bellek kullanımını minimize ederek sistemin kararlı çalışmasını "
        "sağlamaktadır.",
        False,
    ),
    (
        "Sonuç olarak üçüncü hafta itibarıyla Restoran Yönetim Sistemi; gerçek zamanlı QR sipariş, "
        "kategori filtreleme, kampanya yönetimi, günlük istatistik raporu ve tam modüler OOP mimarisi "
        "ile işlevsel ve zengin özellikli bir uygulama haline getirilmiştir.",
        False,
    ),
]

for metin, baslik in PARAGRAFLAR:
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    run = p.add_run(metin)
    run.font.size = Pt(11)
    if baslik:
        run.bold = True
        p.paragraph_format.space_before = Pt(10)
    else:
        p.paragraph_format.first_line_indent = Cm(1)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

# ── İMZA TABLOSU ────────────────────────────────────────────────
doc.add_page_break()

it = doc.add_table(rows=4, cols=3)
it.style = "Table Grid"
it.alignment = WD_TABLE_ALIGNMENT.CENTER

for i, b in enumerate(["NUMARA", "GRUP ÜYELERİNİN İSİMLERİ", "İMZA"]):
    cell_font(it.rows[0].cells[i], b, bold=True, size=10,
              align=WD_ALIGN_PARAGRAPH.CENTER)
    set_cell_bg(it.rows[0].cells[i], "2E75B6")
    it.rows[0].cells[i].paragraphs[0].runs[0].font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)

uyeler = [
    ("202407124023", "Ege Recep Alembeyli"),
    ("202407124030", "Taha Saranbeyli"),
]
for idx, (numara, isim) in enumerate(uyeler, 1):
    cell_font(it.rows[idx].cells[0], numara, size=10, align=WD_ALIGN_PARAGRAPH.CENTER)
    cell_font(it.rows[idx].cells[1], isim,   size=10)
    it.rows[idx].cells[2].text = ""

cell_font(it.rows[3].cells[0], "", size=10)
cell_font(it.rows[3].cells[1], "Öğr.Gör. İlker Duran", bold=True, size=10)
cell_font(it.rows[3].cells[2], "İmza", size=10, align=WD_ALIGN_PARAGRAPH.CENTER)

kayit = r"C:\Users\Doyuum\Desktop\BLP_Grup11_3Hafta_Rapor.docx"
doc.save(kayit)
print("Kaydedildi:", kayit)
